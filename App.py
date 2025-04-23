<<<<<<< HEAD
#--------------- NEW Design CODE-------------------------#
import streamlit as st
import google.generativeai as genai
import fitz
import docx
import pandas as pd

# --- PAGE CONFIG ---
st.set_page_config(page_title="Tech Career Coach", layout="wide")

# --- SIDEBAR NAVIGATION ---
st.sidebar.title("🔍 Navigation")
menu = st.sidebar.radio("Go to", ["Chatbot", "Resume Review", "Mentor Match", "Learning Path"])

# --- GEMINI CONFIG ---
api_key = st.secrets["GEMINI_API_KEY"]
genai.configure(api_key=api_key)
model = genai.GenerativeModel("gemini-1.5-flash")

# --- SESSION STATE ---
if "messages" not in st.session_state:
    st.session_state["messages"] = []

# --- PROMPT FUNCTION ---
def ask_single_prompt(prompt_text):
    try:
        response = model.generate_content(prompt_text)
        return response.text
    except Exception as e:
        st.error(f"❌ Error: {str(e)}")
        return None

# --- CHATBOT PAGE ---
if menu == "Chatbot":
    st.title("💬 Tech Career Coach Chatbot")
    st.caption("Ask about careers, mentorship, or learning paths.")

    for msg in st.session_state["messages"]:
        who = "🧑" if msg["role"] == "user" else "🤖"
        st.markdown(f"{who}: **{msg['parts']}**")

    user_input = st.chat_input("Type your question...")
    if user_input:
        st.session_state["messages"].append({"role": "user", "parts": user_input})
        try:
            response = model.generate_content(st.session_state["messages"])
            st.session_state["messages"].append({"role": "model", "parts": response.text})
            st.rerun()
        except Exception as e:
            st.error(f"❌ Error: {str(e)}")

# --- RESUME REVIEW PAGE ---
elif menu == "Resume Review":
    st.title("📎 Resume Reviewer")
    uploaded_file = st.file_uploader("Upload your resume (PDF, DOCX, or TXT)", type=["pdf", "docx", "txt"])

    def extract_text_from_file(uploaded_file):
        file_type = uploaded_file.name.split('.')[-1].lower()
        if file_type == "txt":
            return uploaded_file.read().decode("utf-8")
        elif file_type == "pdf":
            pdf = fitz.open(stream=uploaded_file.read(), filetype="pdf")
            return "".join([page.get_text() for page in pdf])
        elif file_type == "docx":
            doc = docx.Document(uploaded_file)
            return "\n".join([para.text for para in doc.paragraphs])
        return "Unsupported file type."

    if uploaded_file:
        extracted_text = extract_text_from_file(uploaded_file)
        st.markdown("**: Uploaded resume for review.**")
        reply = ask_single_prompt(f"Please review this resume:{extracted_text}")
        if reply:
            st.markdown(f"🤖: {reply}")

# --- MENTOR MATCH PAGE ---
elif menu == "Mentor Match":
    st.title("Mentor Match")
    st.markdown("Find mentors based on your interests, preferences, and location.")

    df = pd.read_csv("mentors.csv")

    with st.container():
        col1, col2, col3 = st.columns(3)
        field = col1.selectbox("🎯 Field of interest", sorted(df["field"].unique()))
        gender = col2.selectbox("🚻 Preferred gender", ["Any", "Female", "Male"])
        location = col3.selectbox("📍 Preferred location", ["Any"] + sorted(df["location"].unique()))

    # Filter logic
    filtered = df[df["field"] == field]
    if gender != "Any":
        filtered = filtered[filtered["gender"] == gender]
    if location != "Any":
        filtered = filtered[filtered["location"] == location]

    st.divider()
    st.subheader("✨ Top Mentor Recommendations")

    # if not filtered.empty:
    #     for _, mentor in filtered.head(3).iterrows():
    #         st.markdown(f"""
    #         <div style='padding: 1rem; border: 1px solid #DDD; border-radius: 10px; margin-bottom: 1rem; background-color: #f9f9f9;'>
    #             <h4 style="margin-bottom: 0.3rem;">{mentor['name']}</h4>
    #             <p style="margin: 0.2rem 0;">
    #                 <strong>📍 Location:</strong> {mentor['location']} &nbsp; | &nbsp;
    #                 <strong>💼 Field:</strong> {mentor['field']} &nbsp; | &nbsp;
    #                 <strong>🚻 Gender:</strong> {mentor['gender']}
    #             </p>
    #             <p style="margin: 0.2rem 0;">
    #                 <strong>📧 Email:</strong> <code>{mentor['email']}</code> &nbsp; | &nbsp;
    #                 <strong>🕓 Experience:</strong> {mentor['experience']} years
    #             </p>
    #         </div>
    #         """, unsafe_allow_html=True)
    #            else: st.warning("No mentors match your filters. Try adjusting the criteria.")

                                 
       
if not filtered.empty:
    for _, mentor in filtered.head(3).iterrows():
        st.markdown(f"""
        <div style='
            padding: 1.2rem;
            border: 1px solid #ccc;
            border-radius: 12px;
            margin-bottom: 1.2rem;
            background-color: #f1f6f9;
            box-shadow: 2px 2px 8px rgba(0,0,0,0.05);
        '>
            <h4 style="
                color: #005f73;
                margin-bottom: 0.5rem;
                font-weight: 700;
                font-size: 1.3rem;
            ">{mentor['name']}</h4>

            <p style="margin: 0.2rem 0;">
                <span style="color: #444;"><strong>📍 Location:</strong></span>
                <span style="color: #000;">{mentor['location']}</span> &nbsp; | &nbsp;

                <span style="color: #444;"><strong>💼 Field:</strong></span>
                <span style="color: #000;">{mentor['field']}</span> &nbsp; | &nbsp;

                <span style="color: #444;"><strong>🚻 Gender:</strong></span>
                <span style="color: #000;">{mentor['gender']}</span>
            </p>

            <p style="margin: 0.2rem 0;">
                <span style="color: #444;"><strong>📧 Email:</strong></span>
                <code style="
                    background-color: #eaf4fc;
                    padding: 2px 6px;
                    border-radius: 4px;
                    font-size: 0.95rem;
                ">{mentor['email']}</code> &nbsp; | &nbsp;

                <span style="color: #444;"><strong>🕓 Experience:</strong></span>
                <span style="color: #000;">{mentor['experience']} years</span>
            </p>
        </div>
        """, unsafe_allow_html=True)
else:
    st.warning("No mentors match your filters. Try adjusting the criteria.")

# --- LEARNING PATH PAGE ---
elif menu == "Learning Path":
    st.title("📚 Personalized Learning Path")
    domain = st.selectbox("Choose a career track:", [
        "Frontend Development", "Data Science", "Cloud Engineering",
        "Cybersecurity", "Product Management", "AI/ML",
        "UX Design", "Backend Development", "Full Stack Development", "Data Engineering"
    ])
    if st.button("Suggest Learning Path"):
        prompt = f"Suggest a beginner-to-intermediate learning path using freeCodeCamp or Coursera for someone interested in becoming a {domain}. Include key skills and timeline."
        reply = ask_single_prompt(prompt)
        if reply:
            st.markdown(f"🤖: {reply}")

=======
#--------------- NEW Design CODE-------------------------#
import streamlit as st
import google.generativeai as genai
import fitz
import docx
import pandas as pd

# --- PAGE CONFIG ---
st.set_page_config(page_title="Tech Career Coach", layout="wide")

# --- SIDEBAR NAVIGATION ---
st.sidebar.title("🔍 Navigation")
menu = st.sidebar.radio("Go to", ["Chatbot", "Resume Review", "Mentor Match", "Learning Path"])

# --- GEMINI CONFIG ---
api_key = st.secrets["GEMINI_API_KEY"]
genai.configure(api_key=api_key)
model = genai.GenerativeModel("gemini-1.5-flash")

# --- SESSION STATE ---
if "messages" not in st.session_state:
    st.session_state["messages"] = []

# --- PROMPT FUNCTION ---
def ask_single_prompt(prompt_text):
    try:
        response = model.generate_content(prompt_text)
        return response.text
    except Exception as e:
        st.error(f"❌ Error: {str(e)}")
        return None

# --- CHATBOT PAGE ---
if menu == "Chatbot":
    st.title("💬 Tech Career Coach Chatbot")
    st.caption("Ask about careers, mentorship, or learning paths.")

    for msg in st.session_state["messages"]:
        who = "🧑" if msg["role"] == "user" else "🤖"
        st.markdown(f"{who}: **{msg['parts']}**")

    user_input = st.chat_input("Type your question...")
    if user_input:
        st.session_state["messages"].append({"role": "user", "parts": user_input})
        try:
            response = model.generate_content(st.session_state["messages"])
            st.session_state["messages"].append({"role": "model", "parts": response.text})
            st.rerun()
        except Exception as e:
            st.error(f"❌ Error: {str(e)}")

# --- RESUME REVIEW PAGE ---
elif menu == "Resume Review":
    st.title("📎 Resume Reviewer")
    uploaded_file = st.file_uploader("Upload your resume (PDF, DOCX, or TXT)", type=["pdf", "docx", "txt"])

    def extract_text_from_file(uploaded_file):
        file_type = uploaded_file.name.split('.')[-1].lower()
        if file_type == "txt":
            return uploaded_file.read().decode("utf-8")
        elif file_type == "pdf":
            pdf = fitz.open(stream=uploaded_file.read(), filetype="pdf")
            return "".join([page.get_text() for page in pdf])
        elif file_type == "docx":
            doc = docx.Document(uploaded_file)
            return "\n".join([para.text for para in doc.paragraphs])
        return "Unsupported file type."

    if uploaded_file:
        extracted_text = extract_text_from_file(uploaded_file)
        st.markdown("**: Uploaded resume for review.**")
        reply = ask_single_prompt(f"Please review this resume:{extracted_text}")
        if reply:
            st.markdown(f"🤖: {reply}")

# --- MENTOR MATCH PAGE ---
elif menu == "Mentor Match":
    st.title("Mentor Match")
    st.markdown("Find mentors based on your interests, preferences, and location.")

    df = pd.read_csv("mentors.csv")

    with st.container():
        col1, col2, col3 = st.columns(3)
        field = col1.selectbox("🎯 Field of interest", sorted(df["field"].unique()))
        gender = col2.selectbox("🚻 Preferred gender", ["Any", "Female", "Male"])
        location = col3.selectbox("📍 Preferred location", ["Any"] + sorted(df["location"].unique()))

    # Filter logic
    filtered = df[df["field"] == field]
    if gender != "Any":
        filtered = filtered[filtered["gender"] == gender]
    if location != "Any":
        filtered = filtered[filtered["location"] == location]

    st.divider()
    st.subheader("✨ Top Mentor Recommendations")

    if not filtered.empty:
        for _, mentor in filtered.head(3).iterrows():
            st.markdown(f"""
            <div style='padding: 1rem; border: 1px solid #DDD; border-radius: 10px; margin-bottom: 1rem; background-color: #f9f9f9;'>
                <h4 style="margin-bottom: 0.3rem;">{mentor['name']}</h4>
                <p style="margin: 0.2rem 0;">
                    <strong>📍 Location:</strong> {mentor['location']} &nbsp; | &nbsp;
                    <strong>💼 Field:</strong> {mentor['field']} &nbsp; | &nbsp;
                    <strong>🚻 Gender:</strong> {mentor['gender']}
                </p>
                <p style="margin: 0.2rem 0;">
                    <strong>📧 Email:</strong> <code>{mentor['email']}</code> &nbsp; | &nbsp;
                    <strong>🕓 Experience:</strong> {mentor['experience']} years
                </p>
            </div>
            """, unsafe_allow_html=True)
    else:
        st.warning("No mentors match your filters. Try adjusting the criteria.")


# --- LEARNING PATH PAGE ---
elif menu == "Learning Path":
    st.title("📚 Personalized Learning Path")
    domain = st.selectbox("Choose a career track:", [
        "Frontend Development", "Data Science", "Cloud Engineering",
        "Cybersecurity", "Product Management", "AI/ML",
        "UX Design", "Backend Development", "Full Stack Development", "Data Engineering"
    ])
    if st.button("Suggest Learning Path"):
        prompt = f"Suggest a beginner-to-intermediate learning path using freeCodeCamp or Coursera for someone interested in becoming a {domain}. Include key skills and timeline."
        reply = ask_single_prompt(prompt)
        if reply:
            st.markdown(f"🤖: {reply}")

#>>>>>>> 9b5e4251fd6cd9ed7b316f8fa84c888b5819e84b
